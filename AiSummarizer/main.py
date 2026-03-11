from io import BytesIO
from os import getenv
from pathlib import Path
import re
from typing import Any

from dotenv import load_dotenv

load_dotenv()

import httpx
from docx import Document
from fastapi import Depends, FastAPI, HTTPException, Request, UploadFile
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field, ValidationError
from pypdf import PdfReader

MAX_DOCUMENT_CHARS = 50_000
MAX_UPLOAD_SIZE_BYTES = 10 * 1024 * 1024
SUPPORTED_UPLOAD_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".rtf"}


class SummarizeRequest(BaseModel):
    document_text: str = Field(..., max_length=MAX_DOCUMENT_CHARS, description="Plain text from the legal document")
    document_title: str | None = Field(default=None, max_length=250, description="Optional document title")
    max_summary_words: int = Field(default=200, ge=50, le=1000, description="Approximate summary length")


class SummarizeResponse(BaseModel):
    summary: str
    model: str


class GeminiAPIError(Exception):
    def __init__(self, status_code: int, detail: str):
        self.status_code = status_code
        self.detail = detail
        super().__init__(detail)


class DocumentProcessingError(Exception):
    def __init__(self, status_code: int, detail: str):
        self.status_code = status_code
        self.detail = detail
        super().__init__(detail)


class GeminiSummarizer:
    def __init__(
        self,
        api_key: str,
        model: str = "gemini-2.5-flash-lite",
        timeout: float = 30.0,
        base_url: str = "https://generativelanguage.googleapis.com/v1beta/models",
    ):
        self.api_key = api_key
        self.model = model
        self.timeout = timeout
        self.base_url = base_url.rstrip("/")

    async def summarize_document(
        self,
        document_text: str,
        document_title: str | None = None,
        max_summary_words: int = 200,
    ) -> str:
        if not self.api_key:
            raise GeminiAPIError(500, "GEMINI_API_KEY is not configured.")

        prompt = (
            "You are a legal-document summarizer. Summarize the following legal document in plain language. "
            f"Keep the summary to about {max_summary_words} words. Highlight the main parties, obligations, dates, risks, and critical terms.\n\n"
            f"Title: {document_title or 'Untitled legal document'}\n\n"
            f"Document:\n{document_text}"
        )

        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": 0.2},
        }

        async with httpx.AsyncClient(timeout=self.timeout) as client:
            try:
                response = await client.post(
                    f"{self.base_url}/{self.model}:generateContent",
                    params={"key": self.api_key},
                    json=payload,
                )
                response.raise_for_status()
            except httpx.TimeoutException as exc:
                raise GeminiAPIError(504, "Gemini API timed out while generating the summary.") from exc
            except httpx.HTTPStatusError as exc:
                raise GeminiAPIError(*self._map_http_error(exc.response)) from exc
            except httpx.HTTPError as exc:
                raise GeminiAPIError(502, "Could not reach the Gemini API.") from exc

        return self._extract_summary(response.json())

    @staticmethod
    def _map_http_error(response: httpx.Response) -> tuple[int, str]:
        message = "Gemini API request failed."
        try:
            error_body = response.json()
            message = error_body.get("error", {}).get("message", message)
        except ValueError:
            pass

        if response.status_code in {401, 403}:
            return 502, "Gemini API rejected the credentials or model access."
        if response.status_code == 429:
            return 429, "Gemini API rate limit reached. Please try again shortly."
        if response.status_code >= 500:
            return 502, "Gemini API is temporarily unavailable."
        return 502, message

    @staticmethod
    def _extract_summary(data: dict[str, Any]) -> str:
        summary_parts: list[str] = []
        for candidate in data.get("candidates", []):
            content = candidate.get("content", {})
            for part in content.get("parts", []):
                text = part.get("text")
                if text:
                    summary_parts.append(text.strip())

        summary = "\n".join(part for part in summary_parts if part).strip()
        if summary:
            return summary

        block_reason = data.get("promptFeedback", {}).get("blockReason")
        if block_reason:
            raise GeminiAPIError(502, f"Gemini blocked the summary request: {block_reason}.")

        raise GeminiAPIError(502, "Gemini API returned an unexpected response format.")


app = FastAPI(title="AI Legal Document Summarizer", version="1.1.0")


@app.exception_handler(GeminiAPIError)
async def gemini_api_error_handler(_: Request, exc: GeminiAPIError) -> JSONResponse:
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


@app.exception_handler(DocumentProcessingError)
async def document_processing_error_handler(_: Request, exc: DocumentProcessingError) -> JSONResponse:
    return JSONResponse(status_code=exc.status_code, content={"detail": exc.detail})


def get_summarizer() -> GeminiSummarizer:
    return GeminiSummarizer(api_key=getenv("GEMINI_API_KEY", "").strip())


def normalize_document_text(text: str) -> str:
    normalized_lines = [line.strip() for line in text.replace("\x00", "").splitlines()]
    return "\n".join(line for line in normalized_lines if line).strip()


def extract_text_from_text_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError(400, "The uploaded text file could not be decoded.")


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # pragma: no cover - library-specific parsing failures
        raise DocumentProcessingError(400, "The uploaded PDF could not be read.") from exc


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:  # pragma: no cover - library-specific parsing failures
        raise DocumentProcessingError(400, "The uploaded Word document could not be read.") from exc

    paragraph_text = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
    return "\n".join(paragraph_text + table_text)


def extract_text_from_rtf_bytes(file_bytes: bytes) -> str:
    raw_text = file_bytes.decode("utf-8", errors="ignore")
    raw_text = raw_text.replace("\\par", "\n")
    raw_text = re.sub(r"\\'[0-9a-fA-F]{2}", "", raw_text)
    raw_text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw_text)
    raw_text = raw_text.replace("{", "").replace("}", "")
    return "\n".join(line.strip() for line in raw_text.splitlines()).strip()


async def extract_text_from_upload(upload: UploadFile) -> str:
    filename = upload.filename or "uploaded-document"
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_UPLOAD_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_UPLOAD_EXTENSIONS))
        raise DocumentProcessingError(415, f"Unsupported file type '{extension or 'unknown'}'. Supported types: {supported}.")

    file_bytes = await upload.read()
    if not file_bytes:
        raise DocumentProcessingError(400, "The uploaded file is empty.")
    if len(file_bytes) > MAX_UPLOAD_SIZE_BYTES:
        raise DocumentProcessingError(413, "The uploaded file exceeds the 10 MB limit.")

    if extension in {".txt", ".md"}:
        document_text = extract_text_from_text_bytes(file_bytes)
    elif extension == ".pdf":
        document_text = extract_text_from_pdf_bytes(file_bytes)
    elif extension == ".docx":
        document_text = extract_text_from_docx_bytes(file_bytes)
    else:
        document_text = extract_text_from_rtf_bytes(file_bytes)

    document_text = normalize_document_text(document_text)
    if not document_text:
        raise DocumentProcessingError(400, "No readable text could be extracted from the uploaded document.")

    return document_text


def build_request_payload(document_text: str, document_title: str | None, max_summary_words: int) -> SummarizeRequest:
    try:
        payload = SummarizeRequest(
            document_text=document_text,
            document_title=document_title,
            max_summary_words=max_summary_words,
        )
    except ValidationError as exc:
        raise HTTPException(status_code=422, detail=exc.errors()) from exc

    payload.document_text = payload.document_text.strip()
    if not payload.document_text:
        raise HTTPException(status_code=400, detail="document_text must not be empty.")
    return payload


async def parse_request_payload(request: Request) -> SummarizeRequest:
    content_type = request.headers.get("content-type", "")

    if content_type.startswith("application/json"):
        try:
            body = await request.json()
        except ValueError as exc:
            raise HTTPException(status_code=400, detail="Request body must contain valid JSON.") from exc

        try:
            payload = SummarizeRequest.model_validate(body)
        except ValidationError as exc:
            raise HTTPException(status_code=422, detail=exc.errors()) from exc

        payload.document_text = payload.document_text.strip()
        if not payload.document_text:
            raise HTTPException(status_code=400, detail="document_text must not be empty.")
        return payload

    if content_type.startswith("multipart/form-data"):
        form = await request.form()
        upload = form.get("file")
        if upload is None or not hasattr(upload, "filename") or not hasattr(upload, "read"):
            raise HTTPException(status_code=400, detail="Multipart requests must include a file field named 'file'.")

        raw_max_summary_words = form.get("max_summary_words", 200)
        try:
            max_summary_words = int(raw_max_summary_words)
        except (TypeError, ValueError) as exc:
            raise HTTPException(status_code=422, detail="max_summary_words must be an integer.") from exc

        document_title = form.get("document_title") or Path(upload.filename or "uploaded-document").stem
        document_text = await extract_text_from_upload(upload)
        return build_request_payload(document_text, document_title, max_summary_words)

    raise HTTPException(
        status_code=415,
        detail="Unsupported content type. Use application/json for raw text or multipart/form-data for file uploads.",
    )


@app.get("/")
async def root() -> dict[str, str]:
    return {"message": "Legal document summarizer API is running."}


@app.get("/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/summarize", response_model=SummarizeResponse)
async def summarize_legal_document(
    request: Request,
    summarizer: GeminiSummarizer = Depends(get_summarizer),
) -> SummarizeResponse:
    payload = await parse_request_payload(request)
    summary = await summarizer.summarize_document(
        document_text=payload.document_text,
        document_title=payload.document_title,
        max_summary_words=payload.max_summary_words,
    )
    return SummarizeResponse(summary=summary, model=summarizer.model)


@app.get("/hello/{name}")
async def say_hello(name: str) -> dict[str, str]:
    return {"message": f"Hello {name}"}
